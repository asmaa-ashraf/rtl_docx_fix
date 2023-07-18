
import docxpy
import docx
import os
print('**Welcome to docx right to left fixer**')
directory = input('Please enter the path of the directory containing the docx files:')
if directory=='' or directory==' ' or directory=='\n':
    directory=os.getcwd()
    print(os.listdir(directory))
def rtl_fix(file_path):
    doc=docx.Document()
    text = docxpy.process(file)
    lines= text.split('\n')
    newlines=''
    for line in lines:
        if newlines=='':
            newlines=u'\u202B' +line
        else:
            newlines=newlines+'\n'+u'\u202B' +line
    t=newlines
    p=doc.add_paragraph( t )
    r = p.add_run()
    font = r.font
    font.complex_script = True
    font.rtl = True
    doc.save(file)
if os.path.isdir(directory):
    print(os.listdir(directory))
    for file in os.listdir(directory):
        if file.endswith('.docx'):
            rtl_fix(file)
else:
    print('sorry this is not a path ')